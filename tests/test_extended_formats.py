"""Integration tests for extended format support through the full pipeline.

Tests registry lookup, content type resolution, and end-to-end pipeline
processing for CSV, JSON, XLSX, and DOCX files.  Follows the same patterns as
``test_file_formats.py`` — mock detectors, pipeline construction via
``object.__new__``, etc.
"""

from __future__ import annotations

import io
import json
from unittest.mock import MagicMock

import pytest

from app.models.extraction import ExtractionResult, SpanMap
from app.models.findings import (
    EntityType,
    Finding,
    RedactionStyle,
    SanitizationLevel,
)
from app.pipeline.orchestrator import (
    FORMAT_REGISTRY,
    FormatHandler,
    SanitizationPipeline,
    _resolve_content_type,
)


# ---------------------------------------------------------------------------
# Format registry tests
# ---------------------------------------------------------------------------


class TestCsvJsonFormatRegistry:
    """Verify the registry maps CSV/JSON MIME types to the correct handlers."""

    def test_csv_in_format_registry(self):
        handler = FORMAT_REGISTRY["text/csv"]
        assert isinstance(handler, FormatHandler)
        assert handler.category == "spreadsheet"

    def test_json_in_format_registry(self):
        handler = FORMAT_REGISTRY["application/json"]
        assert isinstance(handler, FormatHandler)
        assert handler.category == "structured"


# ---------------------------------------------------------------------------
# Content type resolution tests
# ---------------------------------------------------------------------------


class TestCsvJsonResolveContentType:
    def test_resolve_content_type_csv(self):
        assert _resolve_content_type(b"data", "data.csv") == "text/csv"

    def test_resolve_content_type_json(self):
        assert _resolve_content_type(b"data", "data.json") == "application/json"


# ---------------------------------------------------------------------------
# CSV pipeline integration tests
# ---------------------------------------------------------------------------


class TestCsvPipelineIntegration:
    """Verify CSV files flow through the full pipeline correctly."""

    def _make_csv_pipeline(self):
        from app.pipeline.extractors.spreadsheet import CsvExtractor
        from app.pipeline.sanitizers.spreadsheet import CsvSanitizer

        pipeline = object.__new__(SanitizationPipeline)
        pipeline._text_detector = MagicMock()
        pipeline._visual_detector = MagicMock()
        pipeline._instance_cache = {
            CsvExtractor: CsvExtractor(),
            CsvSanitizer: CsvSanitizer(),
        }
        return pipeline

    def test_csv_pipeline_end_to_end(self):
        pipeline = self._make_csv_pipeline()
        content = "Name,Email\nJuan Garcia,juan@test.com\n".encode("utf-8")
        findings = [
            Finding(entity_type=EntityType.PERSON_NAME, original_text="Juan Garcia", score=0.9),
            Finding(entity_type=EntityType.EMAIL, original_text="juan@test.com", score=0.99),
        ]
        pipeline._text_detector.detect.return_value = findings
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.csv",
            level=SanitizationLevel.STANDARD,
        )

        assert result.summary.total_findings == 2
        assert result.sanitized_content is not None
        sanitized_text = result.sanitized_content.decode("utf-8")
        assert "Juan Garcia" not in sanitized_text
        assert "juan@test.com" not in sanitized_text

    def test_csv_output_filename(self):
        pipeline = self._make_csv_pipeline()
        content = "Name\nJuan\n".encode("utf-8")
        pipeline._text_detector.detect.return_value = []
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.csv",
            level=SanitizationLevel.STANDARD,
        )

        assert result.output_filename == "data_sanitized.csv"


# ---------------------------------------------------------------------------
# JSON pipeline integration tests
# ---------------------------------------------------------------------------


class TestJsonPipelineIntegration:
    """Verify JSON files flow through the full pipeline correctly."""

    def _make_json_pipeline(self):
        from app.pipeline.extractors.structured import JsonExtractor
        from app.pipeline.sanitizers.structured import JsonSanitizer

        pipeline = object.__new__(SanitizationPipeline)
        pipeline._text_detector = MagicMock()
        pipeline._visual_detector = MagicMock()
        pipeline._instance_cache = {
            JsonExtractor: JsonExtractor(),
            JsonSanitizer: JsonSanitizer(),
        }
        return pipeline

    def test_json_pipeline_end_to_end(self):
        pipeline = self._make_json_pipeline()
        data = {"name": "Juan Garcia", "email": "juan@test.com"}
        content = json.dumps(data).encode("utf-8")
        findings = [
            Finding(entity_type=EntityType.PERSON_NAME, original_text="Juan Garcia", score=0.9),
            Finding(entity_type=EntityType.EMAIL, original_text="juan@test.com", score=0.99),
        ]
        pipeline._text_detector.detect.return_value = findings
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.json",
            level=SanitizationLevel.STANDARD,
        )

        assert result.summary.total_findings == 2
        assert result.sanitized_content is not None
        output = json.loads(result.sanitized_content)
        assert "Juan Garcia" not in output["name"]
        assert "juan@test.com" not in output["email"]

    def test_json_output_filename(self):
        pipeline = self._make_json_pipeline()
        data = {"city": "Madrid"}
        content = json.dumps(data).encode("utf-8")
        pipeline._text_detector.detect.return_value = []
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.json",
            level=SanitizationLevel.STANDARD,
        )

        assert result.output_filename == "data_sanitized.json"


# ---------------------------------------------------------------------------
# XLSX format registry tests
# ---------------------------------------------------------------------------


class TestXlsxDocxFormatRegistry:
    """Verify the registry maps XLSX/DOCX MIME types to the correct handlers."""

    def test_xlsx_in_format_registry(self):
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        handler = FORMAT_REGISTRY[mime]
        assert isinstance(handler, FormatHandler)
        assert handler.category == "spreadsheet"

    def test_docx_in_format_registry(self):
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        handler = FORMAT_REGISTRY[mime]
        assert isinstance(handler, FormatHandler)
        assert handler.category == "document"


# ---------------------------------------------------------------------------
# XLSX / DOCX content type resolution tests
# ---------------------------------------------------------------------------


class TestXlsxDocxResolveContentType:
    def test_resolve_content_type_xlsx(self):
        assert (
            _resolve_content_type(b"data", "data.xlsx")
            == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    def test_resolve_content_type_docx(self):
        assert (
            _resolve_content_type(b"data", "data.docx")
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


# ---------------------------------------------------------------------------
# XLSX pipeline integration tests
# ---------------------------------------------------------------------------


class TestXlsxPipelineIntegration:
    """Verify XLSX files flow through the full pipeline correctly."""

    def _make_xlsx_pipeline(self):
        from app.pipeline.extractors.spreadsheet import XlsxExtractor
        from app.pipeline.sanitizers.spreadsheet import XlsxSanitizer

        pipeline = object.__new__(SanitizationPipeline)
        pipeline._text_detector = MagicMock()
        pipeline._visual_detector = MagicMock()
        pipeline._instance_cache = {
            XlsxExtractor: XlsxExtractor(),
            XlsxSanitizer: XlsxSanitizer(),
        }
        return pipeline

    @staticmethod
    def _make_xlsx_bytes() -> bytes:
        """Create a minimal XLSX workbook with PII data."""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["Name", "Email"])
        ws.append(["Juan Garcia", "juan@test.com"])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_xlsx_pipeline_end_to_end(self):
        pipeline = self._make_xlsx_pipeline()
        content = self._make_xlsx_bytes()
        findings = [
            Finding(entity_type=EntityType.PERSON_NAME, original_text="Juan Garcia", score=0.9),
            Finding(entity_type=EntityType.EMAIL, original_text="juan@test.com", score=0.99),
        ]
        pipeline._text_detector.detect.return_value = findings
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.xlsx",
            level=SanitizationLevel.STANDARD,
        )

        assert result.summary.total_findings == 2
        assert result.sanitized_content is not None

        # Re-parse the sanitized XLSX and verify PII is gone.
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(result.sanitized_content))
        ws = wb.active
        all_values = " ".join(
            str(cell.value) for row in ws.iter_rows() for cell in row if cell.value
        )
        assert "Juan Garcia" not in all_values
        assert "juan@test.com" not in all_values

    def test_xlsx_output_filename(self):
        pipeline = self._make_xlsx_pipeline()
        content = self._make_xlsx_bytes()
        pipeline._text_detector.detect.return_value = []
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.xlsx",
            level=SanitizationLevel.STANDARD,
        )

        assert result.output_filename == "data_sanitized.xlsx"


# ---------------------------------------------------------------------------
# DOCX pipeline integration tests
# ---------------------------------------------------------------------------


class TestDocxPipelineIntegration:
    """Verify DOCX files flow through the full pipeline correctly."""

    def _make_docx_pipeline(self):
        from app.pipeline.extractors.document import DocxExtractor
        from app.pipeline.sanitizers.document import DocxSanitizer

        pipeline = object.__new__(SanitizationPipeline)
        pipeline._text_detector = MagicMock()
        pipeline._visual_detector = MagicMock()
        pipeline._instance_cache = {
            DocxExtractor: DocxExtractor(),
            DocxSanitizer: DocxSanitizer(),
        }
        return pipeline

    @staticmethod
    def _make_docx_bytes() -> bytes:
        """Create a minimal DOCX document with PII data."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("My name is Juan Garcia and my email is juan@test.com")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_docx_pipeline_end_to_end(self):
        pipeline = self._make_docx_pipeline()
        content = self._make_docx_bytes()
        findings = [
            Finding(entity_type=EntityType.PERSON_NAME, original_text="Juan Garcia", score=0.9),
            Finding(entity_type=EntityType.EMAIL, original_text="juan@test.com", score=0.99),
        ]
        pipeline._text_detector.detect.return_value = findings
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.docx",
            level=SanitizationLevel.STANDARD,
        )

        assert result.summary.total_findings == 2
        assert result.sanitized_content is not None

        # Re-parse the sanitized DOCX and verify PII is gone.
        from docx import Document

        doc = Document(io.BytesIO(result.sanitized_content))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Juan Garcia" not in all_text
        assert "juan@test.com" not in all_text

    def test_docx_output_filename(self):
        pipeline = self._make_docx_pipeline()
        content = self._make_docx_bytes()
        pipeline._text_detector.detect.return_value = []
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.docx",
            level=SanitizationLevel.STANDARD,
        )

        assert result.output_filename == "data_sanitized.docx"


# ---------------------------------------------------------------------------
# HTML / RTF format registry tests
# ---------------------------------------------------------------------------


class TestHtmlRtfFormatRegistry:
    """Verify the registry maps HTML/RTF MIME types to the correct handlers."""

    def test_html_in_format_registry(self):
        handler = FORMAT_REGISTRY["text/html"]
        assert isinstance(handler, FormatHandler)
        assert handler.category == "structured"

    def test_rtf_in_format_registry(self):
        handler = FORMAT_REGISTRY["text/rtf"]
        assert isinstance(handler, FormatHandler)
        assert handler.category == "document"


# ---------------------------------------------------------------------------
# HTML / RTF content type resolution tests
# ---------------------------------------------------------------------------


class TestHtmlRtfResolveContentType:
    def test_resolve_content_type_html(self):
        assert _resolve_content_type(b"data", "page.html") == "text/html"

    def test_resolve_content_type_rtf(self):
        # Python's mimetypes maps .rtf to application/rtf
        assert _resolve_content_type(b"data", "document.rtf") == "application/rtf"


# ---------------------------------------------------------------------------
# HTML pipeline integration tests
# ---------------------------------------------------------------------------


class TestHtmlPipelineIntegration:
    """Verify HTML files flow through the full pipeline correctly."""

    def _make_html_pipeline(self):
        from app.pipeline.extractors.structured import HtmlExtractor
        from app.pipeline.sanitizers.structured import HtmlSanitizer

        pipeline = object.__new__(SanitizationPipeline)
        pipeline._text_detector = MagicMock()
        pipeline._visual_detector = MagicMock()
        pipeline._instance_cache = {
            HtmlExtractor: HtmlExtractor(),
            HtmlSanitizer: HtmlSanitizer(),
        }
        return pipeline

    def test_html_pipeline_end_to_end(self):
        pipeline = self._make_html_pipeline()
        content = b"<html><body><p>Juan Garcia</p></body></html>"
        findings = [
            Finding(entity_type=EntityType.PERSON_NAME, original_text="Juan Garcia", score=0.9),
        ]
        pipeline._text_detector.detect.return_value = findings
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="page.html",
            level=SanitizationLevel.STANDARD,
        )

        assert result.summary.total_findings == 1
        assert result.sanitized_content is not None
        sanitized_text = result.sanitized_content.decode("utf-8")
        assert "Juan Garcia" not in sanitized_text

    def test_html_output_filename(self):
        pipeline = self._make_html_pipeline()
        content = b"<html><body><p>Hello</p></body></html>"
        pipeline._text_detector.detect.return_value = []
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="page.html",
            level=SanitizationLevel.STANDARD,
        )

        assert result.output_filename == "page_sanitized.html"


# ---------------------------------------------------------------------------
# RTF pipeline integration tests
# ---------------------------------------------------------------------------


class TestRtfPipelineIntegration:
    """Verify RTF files flow through the full pipeline correctly."""

    def _make_rtf_pipeline(self):
        from app.pipeline.extractors.document import RtfExtractor
        from app.pipeline.sanitizers.document import RtfSanitizer

        pipeline = object.__new__(SanitizationPipeline)
        pipeline._text_detector = MagicMock()
        pipeline._visual_detector = MagicMock()
        pipeline._instance_cache = {
            RtfExtractor: RtfExtractor(),
            RtfSanitizer: RtfSanitizer(),
        }
        return pipeline

    def test_rtf_pipeline_end_to_end(self):
        pipeline = self._make_rtf_pipeline()
        content = rb"{\rtf1\ansi Hello Juan Garcia\par}"
        findings = [
            Finding(entity_type=EntityType.PERSON_NAME, original_text="Juan Garcia", score=0.9),
        ]
        pipeline._text_detector.detect.return_value = findings
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="document.rtf",
            level=SanitizationLevel.STANDARD,
        )

        assert result.summary.total_findings == 1
        assert result.sanitized_content is not None
        sanitized_text = result.sanitized_content.decode("utf-8")
        assert "Juan Garcia" not in sanitized_text

    def test_rtf_output_filename(self):
        pipeline = self._make_rtf_pipeline()
        content = rb"{\rtf1\ansi Hello world\par}"
        pipeline._text_detector.detect.return_value = []
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="document.rtf",
            level=SanitizationLevel.STANDARD,
        )

        assert result.output_filename == "document_sanitized.rtf"


# ---------------------------------------------------------------------------
# ODS / ODT format registry tests
# ---------------------------------------------------------------------------


class TestOdsOdtFormatRegistry:
    """Verify the registry maps ODS/ODT MIME types to the correct handlers."""

    def test_ods_in_format_registry(self):
        mime = "application/vnd.oasis.opendocument.spreadsheet"
        handler = FORMAT_REGISTRY[mime]
        assert isinstance(handler, FormatHandler)
        assert handler.category == "spreadsheet"

    def test_odt_in_format_registry(self):
        mime = "application/vnd.oasis.opendocument.text"
        handler = FORMAT_REGISTRY[mime]
        assert isinstance(handler, FormatHandler)
        assert handler.category == "document"


# ---------------------------------------------------------------------------
# ODS / ODT content type resolution tests
# ---------------------------------------------------------------------------


class TestOdsOdtResolveContentType:
    def test_resolve_content_type_ods(self):
        assert (
            _resolve_content_type(b"data", "data.ods")
            == "application/vnd.oasis.opendocument.spreadsheet"
        )

    def test_resolve_content_type_odt(self):
        assert (
            _resolve_content_type(b"data", "document.odt")
            == "application/vnd.oasis.opendocument.text"
        )


# ---------------------------------------------------------------------------
# ODS pipeline integration tests
# ---------------------------------------------------------------------------


class TestOdsPipelineIntegration:
    """Verify ODS files flow through the full pipeline correctly."""

    def _make_ods_pipeline(self):
        from app.pipeline.extractors.spreadsheet import OdsExtractor
        from app.pipeline.sanitizers.spreadsheet import OdsSanitizer

        pipeline = object.__new__(SanitizationPipeline)
        pipeline._text_detector = MagicMock()
        pipeline._visual_detector = MagicMock()
        pipeline._instance_cache = {
            OdsExtractor: OdsExtractor(),
            OdsSanitizer: OdsSanitizer(),
        }
        return pipeline

    @staticmethod
    def _make_ods_bytes() -> bytes:
        """Create a minimal ODS spreadsheet with PII data."""
        from odf.opendocument import OpenDocumentSpreadsheet
        from odf.table import Table, TableRow, TableCell
        from odf.text import P

        doc = OpenDocumentSpreadsheet()
        table = Table(name="Sheet1")
        # header row
        row = TableRow()
        cell = TableCell()
        p = P()
        p.addText("Name")
        cell.addElement(p)
        row.addElement(cell)
        table.addElement(row)
        # data row
        row = TableRow()
        cell = TableCell()
        p = P()
        p.addText("Juan Garcia")
        cell.addElement(p)
        row.addElement(cell)
        table.addElement(row)
        doc.spreadsheet.addElement(table)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_ods_pipeline_end_to_end(self):
        pipeline = self._make_ods_pipeline()
        content = self._make_ods_bytes()
        findings = [
            Finding(entity_type=EntityType.PERSON_NAME, original_text="Juan Garcia", score=0.9),
        ]
        pipeline._text_detector.detect.return_value = findings
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.ods",
            level=SanitizationLevel.STANDARD,
        )

        assert result.summary.total_findings == 1
        assert result.sanitized_content is not None

        # Re-parse the sanitized ODS and verify PII is gone.
        from odf.opendocument import load as odf_load
        from odf.text import P as OdfP

        doc = odf_load(io.BytesIO(result.sanitized_content))
        all_text = " ".join(
            str(p) for p in doc.getElementsByType(OdfP)
        )
        assert "Juan Garcia" not in all_text

    def test_ods_output_filename(self):
        pipeline = self._make_ods_pipeline()
        content = self._make_ods_bytes()
        pipeline._text_detector.detect.return_value = []
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="data.ods",
            level=SanitizationLevel.STANDARD,
        )

        assert result.output_filename == "data_sanitized.ods"


# ---------------------------------------------------------------------------
# ODT pipeline integration tests
# ---------------------------------------------------------------------------


class TestOdtPipelineIntegration:
    """Verify ODT files flow through the full pipeline correctly."""

    def _make_odt_pipeline(self):
        from app.pipeline.extractors.document import OdtExtractor
        from app.pipeline.sanitizers.document import OdtSanitizer

        pipeline = object.__new__(SanitizationPipeline)
        pipeline._text_detector = MagicMock()
        pipeline._visual_detector = MagicMock()
        pipeline._instance_cache = {
            OdtExtractor: OdtExtractor(),
            OdtSanitizer: OdtSanitizer(),
        }
        return pipeline

    @staticmethod
    def _make_odt_bytes() -> bytes:
        """Create a minimal ODT document with PII data."""
        from odf.opendocument import OpenDocumentText
        from odf.text import P

        doc = OpenDocumentText()
        p = P()
        p.addText("Contact Juan Garcia")
        doc.text.addElement(p)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_odt_pipeline_end_to_end(self):
        pipeline = self._make_odt_pipeline()
        content = self._make_odt_bytes()
        findings = [
            Finding(entity_type=EntityType.PERSON_NAME, original_text="Juan Garcia", score=0.9),
        ]
        pipeline._text_detector.detect.return_value = findings
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="document.odt",
            level=SanitizationLevel.STANDARD,
        )

        assert result.summary.total_findings == 1
        assert result.sanitized_content is not None

        # Re-parse the sanitized ODT and verify PII is gone.
        from odf.opendocument import load as odf_load
        from odf.text import P as OdfP

        doc = odf_load(io.BytesIO(result.sanitized_content))
        all_text = " ".join(
            str(p) for p in doc.getElementsByType(OdfP)
        )
        assert "Juan Garcia" not in all_text

    def test_odt_output_filename(self):
        pipeline = self._make_odt_pipeline()
        content = self._make_odt_bytes()
        pipeline._text_detector.detect.return_value = []
        pipeline._visual_detector.detect.return_value = []

        result = pipeline.process(
            file_content=content,
            filename="document.odt",
            level=SanitizationLevel.STANDARD,
        )

        assert result.output_filename == "document_sanitized.odt"


# ---------------------------------------------------------------------------
# Registry completeness tests
# ---------------------------------------------------------------------------


class TestRegistryCompleteness:
    """Verify FORMAT_REGISTRY covers all supported MIME types."""

    def test_all_supported_formats_have_registry_entries(self):
        """Every MIME type in settings.SUPPORTED_FORMATS must have a FORMAT_REGISTRY entry."""
        from app.config import settings

        missing = [
            mime for mime in settings.SUPPORTED_FORMATS if mime not in FORMAT_REGISTRY
        ]
        assert missing == [], f"MIME types missing from FORMAT_REGISTRY: {missing}"

    def test_registry_count(self):
        """FORMAT_REGISTRY should have at least 16 entries (8 original + 8 new)."""
        assert len(FORMAT_REGISTRY) >= 16, (
            f"Expected at least 16 registry entries, got {len(FORMAT_REGISTRY)}"
        )
