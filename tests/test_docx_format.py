"""Tests for DOCX format support: DocxExtractor and DocxSanitizer.

Creates fixture documents in-memory using python-docx, then verifies
extraction and sanitization behaviour including cross-run PII replacement.
"""

from __future__ import annotations

from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.shared import Pt

from app.models.extraction import ExtractionResult
from app.models.findings import EntityType, Finding, RedactionStyle
from app.pipeline.extractors.document import DocxExtractor
from app.pipeline.sanitizers.document import DocxSanitizer


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx(*paragraphs: str) -> bytes:
    """Create minimal DOCX bytes with the given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table(paragraphs: list[str], table_data: list[list[str]]) -> bytes:
    """Create DOCX bytes with paragraphs and a table."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_data:
        rows, cols = len(table_data), len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        for r_idx, row_data in enumerate(table_data):
            for c_idx, cell_text in enumerate(row_data):
                table.rows[r_idx].cells[c_idx].text = cell_text
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_cross_run(prefix: str, suffix: str) -> bytes:
    """Create DOCX where a paragraph has text split across two runs.

    This simulates Word splitting "prefix + suffix" across formatting
    boundaries (e.g., bold applied to suffix only).
    """
    doc = Document()
    paragraph = doc.add_paragraph()
    run1 = paragraph.add_run(prefix)
    run2 = paragraph.add_run(suffix)
    run2.bold = True
    run2.font.size = Pt(14)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DocxExtractor tests
# ---------------------------------------------------------------------------


class TestDocxExtractor:
    """Tests for DocxExtractor.extract()."""

    def test_basic_paragraphs(self):
        content = _make_docx("Hello World", "Second paragraph")
        result = DocxExtractor().extract(content, "test.docx")

        assert isinstance(result, ExtractionResult)
        assert "Hello World" in result.text
        assert "Second paragraph" in result.text
        assert result.pages == 1
        assert result.is_scanned is False
        assert len(result.span_map) == 0

    def test_with_tables(self):
        content = _make_docx_with_table(
            paragraphs=["Intro"],
            table_data=[
                ["Name", "Email"],
                ["Juan Garcia", "juan@test.com"],
            ],
        )
        result = DocxExtractor().extract(content, "test.docx")

        assert "Intro" in result.text
        assert "Juan Garcia" in result.text
        assert "juan@test.com" in result.text

    def test_empty_document(self):
        content = _make_docx()
        result = DocxExtractor().extract(content, "empty.docx")

        assert result.text == ""
        assert result.pages == 1
        assert result.is_scanned is False


# ---------------------------------------------------------------------------
# DocxSanitizer tests
# ---------------------------------------------------------------------------


class TestDocxSanitizer:
    """Tests for DocxSanitizer.sanitize()."""

    def test_single_run_pii_black_style(self):
        content = _make_docx("Contact Juan Garcia for details.")
        finding = Finding(
            entity_type=EntityType.PERSON_NAME,
            original_text="Juan Garcia",
            score=0.95,
        )

        sanitized = DocxSanitizer().sanitize(
            content, [finding], "test.docx", style=RedactionStyle.BLACK,
        )

        # Re-parse and verify PII is gone.
        doc = Document(BytesIO(sanitized))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Juan Garcia" not in full_text
        assert "\u2588" * len("Juan Garcia") in full_text

    def test_single_run_pii_placeholder_style(self):
        content = _make_docx("Email: juan@test.com")
        finding = Finding(
            entity_type=EntityType.EMAIL,
            original_text="juan@test.com",
            score=0.99,
        )

        sanitized = DocxSanitizer().sanitize(
            content, [finding], "test.docx", style=RedactionStyle.PLACEHOLDER,
        )

        doc = Document(BytesIO(sanitized))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "juan@test.com" not in full_text
        assert "[EMAIL]" in full_text

    def test_cross_run_pii(self):
        """PII text split across two runs (e.g., bold applied mid-name)."""
        content = _make_docx_cross_run("Juan ", "Garcia")
        finding = Finding(
            entity_type=EntityType.PERSON_NAME,
            original_text="Juan Garcia",
            score=0.95,
        )

        sanitized = DocxSanitizer().sanitize(
            content, [finding], "test.docx", style=RedactionStyle.PLACEHOLDER,
        )

        doc = Document(BytesIO(sanitized))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Juan Garcia" not in full_text
        assert "[PERSON_NAME]" in full_text

    def test_table_cell_pii(self):
        content = _make_docx_with_table(
            paragraphs=[],
            table_data=[
                ["Name", "DNI"],
                ["Juan Garcia", "12345678Z"],
            ],
        )
        findings = [
            Finding(entity_type=EntityType.PERSON_NAME, original_text="Juan Garcia", score=0.9),
            Finding(entity_type=EntityType.DNI_NIE, original_text="12345678Z", score=0.95),
        ]

        sanitized = DocxSanitizer().sanitize(
            content, findings, "test.docx", style=RedactionStyle.PLACEHOLDER,
        )

        doc = Document(BytesIO(sanitized))
        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        assert "Juan Garcia" not in all_text
        assert "12345678Z" not in all_text
        assert "[PERSON_NAME]" in all_text
        assert "[DNI_NIE]" in all_text

    def test_blur_style_uses_entity_label(self):
        content = _make_docx("Phone: 612345678")
        finding = Finding(
            entity_type=EntityType.PHONE,
            original_text="612345678",
            score=0.9,
        )

        sanitized = DocxSanitizer().sanitize(
            content, [finding], "test.docx", style=RedactionStyle.BLUR,
        )

        doc = Document(BytesIO(sanitized))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "612345678" not in full_text
        assert "[PHONE]" in full_text

    def test_no_findings_returns_valid_docx(self):
        content = _make_docx("Nothing to redact here.")

        sanitized = DocxSanitizer().sanitize(
            content, [], "test.docx", style=RedactionStyle.BLACK,
        )

        doc = Document(BytesIO(sanitized))
        assert doc.paragraphs[0].text == "Nothing to redact here."


# ---------------------------------------------------------------------------
# Import guard tests
# ---------------------------------------------------------------------------


class TestImportGuard:
    """Verify import guard raises RuntimeError with install hint."""

    def test_extractor_import_guard(self, monkeypatch):
        import builtins

        import app.pipeline.extractors.document as extractor_mod

        real_import = builtins.__import__

        def mock_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError("No module named 'docx'")
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", mock_import)

        with pytest.raises(RuntimeError, match="python-docx is required"):
            extractor_mod._require_python_docx()

    def test_sanitizer_import_guard(self, monkeypatch):
        import builtins

        import app.pipeline.sanitizers.document as sanitizer_mod

        real_import = builtins.__import__

        def mock_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError("No module named 'docx'")
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", mock_import)

        with pytest.raises(RuntimeError, match="python-docx is required"):
            sanitizer_mod._require_python_docx()
