"""Document extractors for DOCX, RTF, and ODT formats.

Extracts text from ``.docx`` files by iterating paragraphs and table cells,
from ``.rtf`` files by stripping RTF control codes, and from ``.odt`` files
by walking ODF text and table elements.  All return an ``ExtractionResult``
with an empty ``SpanMap`` (no coordinate-based redaction).
"""

from __future__ import annotations

from io import BytesIO

from app.models.extraction import ExtractionResult, SpanMap


def _require_python_docx():
    """Raise a helpful error when *python-docx* is not installed."""
    try:
        import docx  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "python-docx is required for DOCX support. "
            "Install it with: pip install 'saniflow[documents]'"
        ) from None


class DocxExtractor:
    """Extract text content from DOCX files.

    Implements the ``Extractor`` protocol.

    Iterates all paragraphs and table cells, concatenating their text
    with newline separators for natural text flow.
    """

    def extract(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Load a DOCX document and return an ``ExtractionResult``.

        Args:
            file_content: Raw bytes of the DOCX file.
            filename: Original filename (unused beyond protocol compliance).

        Returns:
            An ``ExtractionResult`` with the concatenated text, an empty
            SpanMap, no images, ``pages=1``, and ``is_scanned=False``.
        """
        _require_python_docx()
        from docx import Document

        doc = Document(BytesIO(file_content))
        parts: list[str] = []

        # Extract paragraph text.
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text:
                parts.append(text)

        # Extract table cell text.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    if text:
                        parts.append(text)

        return ExtractionResult(
            text="\n".join(parts),
            span_map=SpanMap(),
            images=[],
            pages=1,
            is_scanned=False,
        )


def _require_striprtf():
    """Raise a helpful error when *striprtf* is not installed."""
    try:
        import striprtf  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "striprtf is required for RTF support. "
            "Install it with: pip install 'saniflow[documents]'"
        ) from None


class RtfExtractor:
    """Extract plain text from RTF files by stripping control codes.

    Implements the ``Extractor`` protocol.

    Uses ``striprtf`` to convert RTF markup to plain text.  Returns an
    ``ExtractionResult`` with an empty ``SpanMap``, ``pages=1``, and
    ``is_scanned=False``.
    """

    def extract(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Decode RTF bytes and return an ``ExtractionResult``.

        Args:
            file_content: Raw bytes of the RTF file.
            filename: Original filename (unused beyond protocol compliance).

        Returns:
            An ``ExtractionResult`` with the stripped plain text, an empty
            SpanMap, no images, ``pages=1``, and ``is_scanned=False``.
        """
        _require_striprtf()
        from striprtf.striprtf import rtf_to_text

        rtf_text = file_content.decode("utf-8", errors="replace")
        plain_text = rtf_to_text(rtf_text)

        return ExtractionResult(
            text=plain_text,
            span_map=SpanMap(),
            images=[],
            pages=1,
            is_scanned=False,
        )


def _require_odfpy():
    """Raise a helpful error when *odfpy* is not installed."""
    try:
        import odf  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "odfpy is required for ODT/ODS support. "
            "Install it with: pip install 'saniflow[documents]'"
        ) from None


class OdtExtractor:
    """Extract text content from ODT (OpenDocument Text) files.

    Implements the ``Extractor`` protocol.

    Iterates all paragraphs and table cells, concatenating their text
    with newline separators for natural text flow.
    """

    def extract(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Load an ODT document and return an ``ExtractionResult``.

        Args:
            file_content: Raw bytes of the ODT file.
            filename: Original filename (unused beyond protocol compliance).

        Returns:
            An ``ExtractionResult`` with the concatenated text, an empty
            SpanMap, no images, ``pages=1``, and ``is_scanned=False``.
        """
        _require_odfpy()
        from odf.opendocument import load
        from odf.table import Table, TableCell, TableRow
        from odf.text import P

        doc = load(BytesIO(file_content))
        parts: list[str] = []

        # Extract paragraph text (top-level body paragraphs).
        for p_elem in doc.text.getElementsByType(P):
            text = "".join(
                t.data for t in p_elem.childNodes if hasattr(t, "data")
            )
            if text:
                parts.append(text)

        # Extract table cell text.
        for table in doc.text.getElementsByType(Table):
            for row in table.getElementsByType(TableRow):
                for cell in row.getElementsByType(TableCell):
                    for p_elem in cell.getElementsByType(P):
                        text = "".join(
                            t.data for t in p_elem.childNodes if hasattr(t, "data")
                        )
                        if text:
                            parts.append(text)

        return ExtractionResult(
            text="\n".join(parts),
            span_map=SpanMap(),
            images=[],
            pages=1,
            is_scanned=False,
        )
