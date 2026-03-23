"""Document sanitizers for DOCX, RTF, and ODT formats.

DOCX: Replaces PII by loading the document, locating
``Finding.original_text`` across paragraphs and table cells, and
replacing with the appropriate redaction string.  Handles cross-run
replacements where Word splits text across formatting runs.

RTF: Strips RTF control codes to plain text, applies find-and-replace
redaction, and returns UTF-8 plain text bytes (not RTF).  This is a
documented limitation — RTF re-encoding is fragile.

ODT: Replaces PII by loading the document via odfpy, walking all
paragraph text nodes in body and tables, and replacing matches
in-place.
"""

from __future__ import annotations

import logging
from io import BytesIO

from app.models.findings import Finding, RedactionStyle

logger = logging.getLogger(__name__)


def _require_python_docx():
    """Raise a helpful error when *python-docx* is not installed."""
    try:
        import docx  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "python-docx is required for DOCX support. "
            "Install it with: pip install 'saniflow[documents]'"
        ) from None


class DocxSanitizer:
    """Redact PII in DOCX files.

    Implements the ``Sanitizer`` protocol.

    Loads the document, iterates paragraphs (including those inside table
    cells), and replaces PII text.  Handles the cross-run problem where
    Word splits a single logical string across multiple formatting runs.
    """

    def sanitize(
        self,
        file_content: bytes,
        findings: list[Finding],
        filename: str,
        *,
        style: RedactionStyle = RedactionStyle.BLACK,
    ) -> bytes:
        """Replace PII text in a DOCX document and return sanitized bytes.

        Args:
            file_content: Original DOCX file bytes.
            findings: PII detections with ``original_text`` populated.
            filename: Original filename (for logging).
            style: Redaction style -- ``BLACK`` uses block characters,
                ``PLACEHOLDER`` / ``BLUR`` use ``[ENTITY_TYPE]`` labels.

        Returns:
            Sanitized DOCX file bytes.
        """
        _require_python_docx()
        from docx import Document

        doc = Document(BytesIO(file_content))

        redaction_count = 0
        for finding in findings:
            if not finding.original_text:
                continue

            replacement = self._get_replacement(finding, style)

            # Process body paragraphs.
            for paragraph in doc.paragraphs:
                if self._replace_in_paragraph(paragraph, finding.original_text, replacement):
                    redaction_count += 1

            # Process table cell paragraphs.
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if self._replace_in_paragraph(
                                paragraph, finding.original_text, replacement,
                            ):
                                redaction_count += 1

        logger.info(
            "%s: redacted %d region(s) across %d finding(s)",
            filename,
            redaction_count,
            len(findings),
        )

        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    # -- Helpers ------------------------------------------------------------

    @staticmethod
    def _replace_in_paragraph(paragraph, original_text: str, replacement: str) -> bool:
        """Replace *original_text* in *paragraph* across runs.

        Word can split a single logical string across multiple formatting
        runs (e.g., "Juan **Garcia**" = 2 runs).  This helper:

        1. Joins all run texts into one string.
        2. Checks if *original_text* is present.
        3. Replaces it and redistributes: ALL new text goes into the first
           run, remaining runs are emptied (preserves document structure).

        Returns ``True`` if a replacement was made.
        """
        runs = paragraph.runs
        if not runs:
            return False

        joined = "".join(run.text for run in runs)
        if original_text not in joined:
            return False

        new_text = joined.replace(original_text, replacement)
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""

        return True

    @staticmethod
    def _get_replacement(finding: Finding, style: RedactionStyle) -> str:
        """Build the replacement string for a given finding and style."""
        if style in (RedactionStyle.PLACEHOLDER, RedactionStyle.BLUR):
            return f"[{finding.entity_type.value}]"
        # BLACK (default): unicode block characters matching original length.
        return "\u2588" * len(finding.original_text)


def _require_striprtf():
    """Raise a helpful error when *striprtf* is not installed."""
    try:
        import striprtf  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "striprtf is required for RTF support. "
            "Install it with: pip install 'saniflow[documents]'"
        ) from None


class RtfSanitizer:
    """Redact PII in RTF files by converting to plain text.

    Implements the ``Sanitizer`` protocol.

    **Design decision**: RTF re-encoding is fragile, so the sanitizer
    outputs UTF-8 plain text, not RTF.  This means RTF input produces
    plain text output (a documented limitation).
    """

    def sanitize(
        self,
        file_content: bytes,
        findings: list[Finding],
        filename: str,
        *,
        style: RedactionStyle = RedactionStyle.BLACK,
    ) -> bytes:
        """Strip RTF markup, replace PII, and return sanitized plain text bytes.

        Args:
            file_content: Original RTF file bytes.
            findings: PII detections with ``original_text`` populated.
            filename: Original filename (for logging).
            style: Redaction style — ``BLACK`` uses block characters,
                ``PLACEHOLDER`` / ``BLUR`` use ``[ENTITY_TYPE]`` labels.

        Returns:
            Sanitized UTF-8 plain text bytes (not RTF).
        """
        _require_striprtf()
        from striprtf.striprtf import rtf_to_text

        rtf_text = file_content.decode("utf-8", errors="replace")
        text = rtf_to_text(rtf_text)

        redaction_count = 0
        for finding in findings:
            if not finding.original_text:
                continue

            pos = text.rfind(finding.original_text)
            if pos == -1:
                continue

            replacement = self._get_replacement(finding, style)
            text = text[:pos] + replacement + text[pos + len(finding.original_text) :]
            redaction_count += 1

        logger.info(
            "%s: redacted %d region(s) across %d finding(s)",
            filename,
            redaction_count,
            len(findings),
        )

        return text.encode("utf-8")

    # -- Helpers ------------------------------------------------------------

    @staticmethod
    def _get_replacement(finding: Finding, style: RedactionStyle) -> str:
        """Build the replacement string for a given finding and style."""
        if style in (RedactionStyle.PLACEHOLDER, RedactionStyle.BLUR):
            return f"[{finding.entity_type.value}]"
        # BLACK (default): unicode block characters matching original length.
        return "\u2588" * len(finding.original_text)


def _require_odfpy():
    """Raise a helpful error when *odfpy* is not installed."""
    try:
        import odf  # noqa: F401
    except ImportError:
        raise RuntimeError(
            "odfpy is required for ODT/ODS support. "
            "Install it with: pip install 'saniflow[documents]'"
        ) from None


class OdtSanitizer:
    """Redact PII in ODT (OpenDocument Text) files.

    Implements the ``Sanitizer`` protocol.

    Loads the document via odfpy, iterates all paragraph text nodes
    (in body and tables), and replaces PII text in-place.
    """

    def sanitize(
        self,
        file_content: bytes,
        findings: list[Finding],
        filename: str,
        *,
        style: RedactionStyle = RedactionStyle.BLACK,
    ) -> bytes:
        """Replace PII text in an ODT document and return sanitized bytes.

        Args:
            file_content: Original ODT file bytes.
            findings: PII detections with ``original_text`` populated.
            filename: Original filename (for logging).
            style: Redaction style -- ``BLACK`` uses block characters,
                ``PLACEHOLDER`` / ``BLUR`` use ``[ENTITY_TYPE]`` labels.

        Returns:
            Sanitized ODT file bytes.
        """
        _require_odfpy()
        from odf.opendocument import load
        from odf.table import Table, TableCell, TableRow
        from odf.text import P

        doc = load(BytesIO(file_content))

        redaction_count = 0
        for finding in findings:
            if not finding.original_text:
                continue

            replacement = self._get_replacement(finding, style)

            # Process all P elements (body paragraphs + table cell paragraphs).
            all_paragraphs = list(doc.text.getElementsByType(P))
            for table in doc.text.getElementsByType(Table):
                for row in table.getElementsByType(TableRow):
                    for cell in row.getElementsByType(TableCell):
                        all_paragraphs.extend(cell.getElementsByType(P))

            for p_elem in all_paragraphs:
                for child in list(p_elem.childNodes):
                    if hasattr(child, "data") and finding.original_text in child.data:
                        child.data = child.data.replace(
                            finding.original_text, replacement, 1,
                        )
                        redaction_count += 1

        logger.info(
            "%s: redacted %d region(s) across %d finding(s)",
            filename,
            redaction_count,
            len(findings),
        )

        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    # -- Helpers ------------------------------------------------------------

    @staticmethod
    def _get_replacement(finding: Finding, style: RedactionStyle) -> str:
        """Build the replacement string for a given finding and style."""
        if style in (RedactionStyle.PLACEHOLDER, RedactionStyle.BLUR):
            return f"[{finding.entity_type.value}]"
        # BLACK (default): unicode block characters matching original length.
        return "\u2588" * len(finding.original_text)
